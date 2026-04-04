#!/usr/bin/env python3

import re
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PARAGRAPH_DIV_CLASSES = {"eyebrow", "meta", "chapter-meta"}


def _clean_text(value: str) -> str:
    value = value.replace("\xa0", " ")
    value = re.sub(r"[ \t]+\n", "\n", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _visible_text(node: Tag) -> str:
    return _clean_text(node.get_text("\n", strip=True))


def _add_multiline_paragraph(
    document: Document,
    text: str,
    *,
    style: str | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    left_indent: float = 0.0,
    italic: bool = False,
) -> None:
    cleaned = _clean_text(text)
    if not cleaned:
        return

    paragraph = document.add_paragraph(style=style)
    if alignment is not None:
        paragraph.alignment = alignment
    if left_indent:
        paragraph.paragraph_format.left_indent = Inches(left_indent)
    paragraph.paragraph_format.space_after = Pt(8)

    lines = cleaned.splitlines()
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.italic = italic
        if index < len(lines) - 1:
            run.add_break()


def _add_rule(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B7B7B7")
    border.append(bottom)
    paragraph._p.get_or_add_pPr().append(border)


def _render_children(document: Document, parent: Tag, state: dict[str, bool]) -> None:
    for child in parent.children:
        if isinstance(child, NavigableString):
            text = _clean_text(str(child))
            if text:
                _add_multiline_paragraph(document, text)
            continue

        if not isinstance(child, Tag):
            continue

        if child.name == "section":
            classes = set(child.get("class", []))
            if "chapter" in classes:
                if state["seen_chapter"]:
                    document.add_page_break()
                state["seen_chapter"] = True
            _render_children(document, child, state)
            if "title-page" in classes:
                document.add_page_break()
            continue

        if child.name == "h1":
            _add_multiline_paragraph(
                document,
                _visible_text(child),
                style="Title",
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )
            continue

        if child.name == "h2":
            document.add_heading(_visible_text(child), level=1)
            continue

        if child.name == "p":
            _add_multiline_paragraph(document, _visible_text(child))
            continue

        if child.name == "blockquote":
            _add_multiline_paragraph(document, _visible_text(child), left_indent=0.4, italic=True)
            continue

        if child.name == "ul":
            for item in child.find_all("li", recursive=False):
                _add_multiline_paragraph(document, _visible_text(item), style="List Bullet")
            continue

        if child.name == "ol":
            for item in child.find_all("li", recursive=False):
                _add_multiline_paragraph(document, _visible_text(item), style="List Number")
            continue

        if child.name == "li":
            _add_multiline_paragraph(document, _visible_text(child), style="List Bullet")
            continue

        if child.name == "hr":
            _add_rule(document)
            continue

        if child.name == "div":
            classes = set(child.get("class", []))
            if classes.intersection(PARAGRAPH_DIV_CLASSES):
                alignment = WD_ALIGN_PARAGRAPH.CENTER if "eyebrow" in classes else None
                _add_multiline_paragraph(document, _visible_text(child), alignment=alignment)
            else:
                _render_children(document, child, state)
            continue

        _render_children(document, child, state)


def convert_html_file_to_docx(
    html_path: str | Path,
    docx_path: str | Path,
    *,
    title: str,
    author: str,
) -> Path:
    input_path = Path(html_path).expanduser().resolve()
    output_path = Path(docx_path).expanduser().resolve()

    soup = BeautifulSoup(input_path.read_text(encoding="utf-8"), "html.parser")
    document = Document()

    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Georgia"
    normal_style.font.size = Pt(11)
    try:
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    except Exception:
        pass

    document.core_properties.title = title
    document.core_properties.author = author

    body = soup.body or soup
    _render_children(document, body, {"seen_chapter": False})
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path
